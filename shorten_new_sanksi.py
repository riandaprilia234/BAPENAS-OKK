import urllib.request
import urllib.parse
import os
import docx

long_url = "https://docs.google.com/forms/d/1illDoXTq-29OGMhTBr4wcbDC7L4siVmxzN-CJposIGc/viewform"
aliases = [
    "FormSanksiOKK2026",
    "FormPengumpulanSanksiOKK26",
    "FormSanksiBapenas2026",
    "SanksiOKK2026"
]

short_url = None
for alias in aliases:
    api_url = f"https://tinyurl.com/api-create.php?url={urllib.parse.quote(long_url)}&alias={alias}"
    try:
        req = urllib.request.Request(api_url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as res:
            short_url = res.read().decode('utf-8')
            print(f"SUCCESS: {short_url} (with alias {alias})")
            break
    except Exception as e:
        print(f"Alias {alias} failed: {e}")

if not short_url:
    # fallback to random short url
    api_url = f"https://tinyurl.com/api-create.php?url={urllib.parse.quote(long_url)}"
    try:
        req = urllib.request.Request(api_url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as res:
            short_url = res.read().decode('utf-8')
            print(f"SUCCESS fallback: {short_url}")
    except Exception as e:
        print("Fallback failed too:", e)

# If we got a short URL, update files
if short_url:
    search_dir = r"C:\Users\LENOVO\Documents\BAPENAS 2026"
    replacements = {
        "tinyurl.com/FormPengumpulanSanksiOKK2026": short_url.replace("https://", ""),
        "bit.ly/FormPengumpulanSanksiOKK2026": short_url.replace("https://", ""),
        "bit.ly/FormPengumpulanSanksiOKK2025": short_url.replace("https://", ""),
        "bit.ly/FormPengumpulanSanksiOKK2024": short_url.replace("https://", ""),
    }

    # Update txt file
    txt_file = os.path.join(search_dir, "LINK PENGUMPULAN SANKSI OKK.txt")
    try:
        with open(txt_file, "w", encoding="utf-8") as f:
            f.write(short_url + "\n")
        print("Updated LINK PENGUMPULAN SANKSI OKK.txt")
    except Exception as e:
        print("Error txt:", e)

    # Update docx files
    for f in os.listdir(search_dir):
        if f.endswith('.docx') and not f.startswith('~$'):
            path = os.path.join(search_dir, f)
            try:
                doc = docx.Document(path)
                changed = False
                for p in doc.paragraphs:
                    for old, new in replacements.items():
                        if old in p.text:
                            p.text = p.text.replace(old, new)
                            changed = True
                
                # tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for old, new in replacements.items():
                                    if old in p.text:
                                        p.text = p.text.replace(old, new)
                                        changed = True
                
                if changed:
                    doc.save(path)
                    print(f"Updated docx: {f}")
            except Exception as e:
                print(f"Error docx {f}: {e}")
