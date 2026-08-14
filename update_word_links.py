import os
import docx

search_dir = r"C:\Users\LENOVO\Documents\BAPENAS 2026"
replacements = {
    "bit.ly/FormPengumpulanSanksiOKK2026": "tinyurl.com/FormPengumpulanSanksiOKK2026",
    "bit.ly/FormPengumpulanSanksiOKK2025": "tinyurl.com/FormPengumpulanSanksiOKK2026",
    "bit.ly/FormPengumpulanSanksiOKK2024": "tinyurl.com/FormPengumpulanSanksiOKK2026",
    "bit.ly/FormPengumpulanSanksiOKK": "tinyurl.com/FormPengumpulanSanksiOKK2026"
}

# Update text file
txt_file = os.path.join(search_dir, "LINK PENGUMPULAN SANKSI OKK.txt")
try:
    with open(txt_file, "w", encoding="utf-8") as f:
        f.write("https://tinyurl.com/FormPengumpulanSanksiOKK2026\n")
    print("Updated LINK PENGUMPULAN SANKSI OKK.txt")
except Exception as e:
    print("Error updating text file:", e)

# Update docx files
for f in os.listdir(search_dir):
    if f.endswith('.docx') and not f.startswith('~$'):
        path = os.path.join(search_dir, f)
        try:
            doc = docx.Document(path)
            changed = False
            for p in doc.paragraphs:
                for old, new in replacements.items():
                    if old in p.text:
                        p.text = p.text.replace(old, new)
                        changed = True
            
            # check inside tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for old, new in replacements.items():
                                if old in p.text:
                                    p.text = p.text.replace(old, new)
                                    changed = True
            
            if changed:
                # Save with the same name (or a temporary copy to avoid lock issues)
                doc.save(path)
                print(f"Updated links in docx: {f}")
        except Exception as e:
            print(f"Error updating docx {f}: {e}")
